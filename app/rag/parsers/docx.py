"""DOCX parser — paragraph text via python-docx (pure Python)."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from app.exceptions.base import RAGDocumentValidationError
from app.rag.parsers.base import ParsedDocument, SourceFormat

# Fixed safety ceilings (parser layer): config may only tighten these.
_MAX_PARAGRAPHS = 200_000
_MAX_TEXT_CHARACTERS = 10_000_000


class DocxParser:
    """Extract paragraph text from a .docx archive.

    Tables are currently not extracted (documented boundary); content
    judgement belongs to the safety layer.  Corrupt or non-zip bytes are
    rejected with a clear message.
    """

    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise RAGDocumentValidationError(
                "无法解析 DOCX 文档：文件损坏或格式不受支持。"
            ) from exc
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        if len(paragraphs) > _MAX_PARAGRAPHS:
            raise RAGDocumentValidationError("DOCX 文档段落数超过限制。")
        text = "\n".join(paragraphs)
        if len(text) > _MAX_TEXT_CHARACTERS:
            raise RAGDocumentValidationError("DOCX 文档内容超过限制。")
        return ParsedDocument(
            filename=filename,
            text=text,
            source_format=_SOURCE_FORMAT,
            page_count=None,
        )


_SOURCE_FORMAT: SourceFormat = "docx"
